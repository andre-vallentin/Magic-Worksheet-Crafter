import logging
from typing import Optional, List, Dict
from docx import Document
from docx.shared import Cm

logger = logging.getLogger(__name__)

class DocumentBuilder:
    def __init__(self, output_path: str, top_margin: float = 0.5, bottom_margin: float = 0.5, left_margin: float = 2.25, right_margin: float = 1.5, colors: Optional[List] = None, logo_path: Optional[str] = None, icon_paths: Optional[Dict] = None):
        self.top_margin = top_margin
        self.bottom_margin = bottom_margin
        self.left_margin = left_margin
        self.right_margin = right_margin
        self.output_path = output_path
        self.logo_path = logo_path
        self.colors = colors
        self.icon_paths = icon_paths or {}
        self.doc = Document()


    def build(self, sections):
        self._setup_margins()

        style = self.doc.styles['Normal']
#        style.font.name = 'Calibri'

        sections['frontmatter'].define_colors(self.colors)
        sections['frontmatter'].define_logo(self.logo_path)
        sections['frontmatter'].build(self.doc)

        self._inject_icon_paths(sections['exercises'])
        self._build_segments(sections['exercises'])

        sections['frontmatter'].build_no_pagenum(self.doc)

        if sections.get('solutions') is not None:
            self.doc.add_heading('LÖSUNGEN', level=2)
            self._inject_icon_paths(sections['solutions'])
            self._build_segments(sections['solutions'])

        if sections.get('references') is not None:
            self.doc.add_heading('QUELLEN', level=2)
            self._build_segments(sections['references'])

    def save(self):
        self.doc.save(self.output_path)


    def _inject_icon_paths(self, segments):
        from .modules.text_task import TextTask
        from .modules.table_task import TableTask
        from .modules.single_choice import SingleChoiceTask
        from .modules.information_text import InformationText

        icon_map = {
            TextTask:         self.icon_paths.get("exercise"),
            TableTask:        self.icon_paths.get("table"),
            SingleChoiceTask: self.icon_paths.get("single_choice"),
            InformationText:  self.icon_paths.get("information"),
        }
        for seg in segments:
            path = icon_map.get(type(seg))
            if path:
                seg.set_icon_path(path)
                
    def _build_segments(self, segments):
        for segment in segments:
            segment.build(self.doc)


    def _setup_margins(self):
        section = self.doc.sections[0]
        section.top_margin    = Cm(self.top_margin)
        section.bottom_margin = Cm(self.bottom_margin)
        section.left_margin   = Cm(self.left_margin)
        section.right_margin  = Cm(self.right_margin)
