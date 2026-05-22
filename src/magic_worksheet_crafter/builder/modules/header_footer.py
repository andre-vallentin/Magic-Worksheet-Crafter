#!/usr/bin/env python3
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .shared.segment import Segment
from .shared.table_column_width import set_column_width
from .study_header import add_study_header


class HeaderFooter(Segment):
    """Manages header and footer content in DOCX documents."""

    def __init__(self, teacher, study, gradeOfSchool, unit, topic):
        self.teacher = teacher
        self.study = study
        self.gradeOfSchool = gradeOfSchool
        self.unit = unit
        self.topic = topic


    def define_colors(self, colors):
        self.colors = colors

    def define_logo(self, logo_path):
        self.logo_path = logo_path

    def build(self, doc):
        self._add_header(doc)
        self._add_footer(doc)

    def build_no_pagenum(self, doc):
        self._add_footer_no_pagenum(doc)

    def _add_header(self, doc):
        """
        Add header table (2 rows × 2 columns).

        Args:
            doc: Document object from python-docx
        """

        section = doc.sections[0]
        header = section.header

        # Clear existing header content
        for paragraph in list(header.paragraphs):
            p = paragraph._element
            p.getparent().remove(p)

        for table in list(header.tables):
            tbl = table._element
            tbl.getparent().remove(tbl)

        usable_width_emu = section.page_width - section.left_margin - section.right_margin
        usable_width_cm = usable_width_emu / 360000

        # Add a 2x2 table inside the header
        table = header.add_table(rows=2, cols=2, width=Cm(usable_width_cm))

        table.autofit = False

        # Add paragraph inside cell
        para = table.cell(0, 0).paragraphs[0]
        
        # Add image to that paragraph
        run = para.add_run()
        run.add_picture(self.logo_path, width=Cm(3.9), height=Cm(0.95))

        add_study_header(para, self.study, self.colors[0], self.colors[1])
        
        table.cell(0, 1).text = self.topic
        table.cell(0, 1).paragraphs[0].style = doc.styles['Heading 1']
        table.cell(0, 1).paragraphs[0].paragraph_format.space_before = Pt(0)
        table.cell(0, 1).paragraphs[0].paragraph_format.space_after = Pt(0)
        table.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in table.cell(0, 1).paragraphs[0].runs:
            run.font.name = 'Calibri'

        table.cell(1, 0).text = 'Datum:'
        for run in table.cell(1, 0).paragraphs[0].runs:
            run.font.name = 'Calibri'

        table.cell(1, 1).text = 'Klasse: ' + self.gradeOfSchool + ' | Fach: ' + self.study + ' | Reihe: ' + self.unit
        table.cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in table.cell(1, 1).paragraphs[0].runs:
            run.font.name = 'Calibri'

        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths using per-row cell approach
        set_column_width(table, 0, 45)    # 4.5 cm = 45 mm
        set_column_width(table, 1, (usable_width_cm - 4.5)*10)

        table.rows[0].height = Cm(1)
        table.rows[1].height = Cm(1)

    def _add_footer(self, doc, start_page=1):
        """
        Add footer with page numbers ONLY for this section.
        """

        section = doc.sections[0]
        footer = section.footer

        # 🔑 IMPORTANT: break link to previous section
        footer.is_linked_to_previous = False

        # Clear existing footer content
        for paragraph in list(footer.paragraphs):
            paragraph._element.getparent().remove(paragraph._element)
        for table in list(footer.tables):
            table._element.getparent().remove(table._element)

        # Info line
        info_para = footer.add_paragraph(
            f"{self.study} | Lehrkraft: {self.teacher} | Klasse: {self.gradeOfSchool} | Reihe: {self.unit}"
        )
        info_para.runs[0].bold = True
        info_para.runs[0].font.size = Pt(10)
        info_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page number line
        page_para = footer.add_paragraph()
        page_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def add_field(paragraph, field_name):
            run = paragraph.add_run()
            fld_begin = OxmlElement('w:fldChar')
            fld_begin.set(qn('w:fldCharType'), 'begin')

            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = f' {field_name} '

            fld_end = OxmlElement('w:fldChar')
            fld_end.set(qn('w:fldCharType'), 'end')

            run._r.append(fld_begin)
            run._r.append(instr)
            run._r.append(fld_end)

        # "Seite X von Y"
        page_para.add_run("Seite ").font.size = Pt(10)
        add_field(page_para, "PAGE")
        page_para.add_run(" von ").font.size = Pt(10)
        add_field(page_para, "SECTIONPAGES")

        # 🔑 Restart page numbering in THIS section
        sectPr = section._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), str(start_page))
        sectPr.append(pgNumType)

    def _add_footer_no_pagenum(self, doc):
        
        section = doc.add_section() # new section for no page numbers in footer
        footer = section.footer
        footer.is_linked_to_previous = False

        # Clear existing footer content
        for paragraph in list(footer.paragraphs):
            p = paragraph._element
            p.getparent().remove(p)
        for table in list(footer.tables):
            tbl = table._element
            tbl.getparent().remove(tbl)

        # Create info paragraph (centered, no page numbers)
        info_para = footer.add_paragraph(
            f"{self.study} | Lehrkraft: {self.teacher} | Klasse: {self.gradeOfSchool} | Reihe: {self.unit}"
        )
        info_para.runs[0].bold = True
        info_para.runs[0].font.size = Pt(10)
        info_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

   