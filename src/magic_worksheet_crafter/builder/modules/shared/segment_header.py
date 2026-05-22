"""
Reusable header table module for creating icon + text headers in Word documents.
Can be used across different classes and document types.
"""

import re
import logging
from pathlib import Path
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .table_column_width import set_column_width

logger = logging.getLogger(__name__)

def segment_header(doc, icon_path, header_text, extra_text='', heading_style='Heading 2',  extra_style='Heading 3',
                     img_width=1.5, col_img_width=2.5, col_text_width=14):
    """
    Add a header table with icon and text to a document.

    Args:
        doc: python-docx Document object
        icon_path: Path to the icon image file
        header_text: Text to display in the header
        heading_style: Heading style to apply (default: 'Heading 3')
        img_width: Image width in cm (default: 1.5)
        col_img_width: Image column width in cm (default: 2.5)
        col_text_width: Text column width in cm (default: 14)

    Returns:
        The created table object
    """
    # Create table with 1 row and 2 columns
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # Set table preferred width (required for column widths to be respected)
    total_width = col_img_width + col_text_width
    tblPr = table._element.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(total_width * 567)))  # cm to dxa
    tblW.set(qn('w:type'), 'dxa')
    tblPr.insert(0, tblW)

    # Configure image cell
    cell_img = table.cell(0, 0)
    run = cell_img.paragraphs[0].add_run()
    if icon_path:
        icon_path_obj = Path(icon_path)
        if icon_path_obj.exists():
            logger.debug(f"Adding icon: {icon_path}")
            run.add_picture(icon_path, width=Cm(img_width))
        else:
            logger.warning(f"Icon path does not exist: {icon_path}")
    else:
        logger.debug(f"No icon path provided for header: {header_text}")

    # Configure text cell
    cell_text = table.cell(0, 1)

    # Set column widths using per-row cell approach
    col_img_width_mm = col_img_width * 10  # Convert cm to mm
    col_text_width_mm = col_text_width * 10  # Convert cm to mm
    set_column_width(table, 0, col_img_width_mm)
    set_column_width(table, 1, col_text_width_mm)

    p_text = cell_text.paragraphs[0]
    p_extra_Text = cell_text.add_paragraph()

    # Parse markdown bold in extra_text
    pattern = r'\*\*(.+?)\*\*'
    last_end = 0

    for match in re.finditer(pattern, extra_text):
        if match.start() > last_end:
            p_extra_Text.add_run(extra_text[last_end:match.start()])

        run = p_extra_Text.add_run(match.group(1))
        run.bold = True
        last_end = match.end()

    if last_end < len(extra_text):
        p_extra_Text.add_run(extra_text[last_end:])

    # Apply heading style
    p_text.style = doc.styles[heading_style]

    p_extra_Text.style = doc.styles[extra_style]

    # Set paragraph formatting
    p_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_text.paragraph_format.left_indent = Cm(0)
    p_text.paragraph_format.first_line_indent = Cm(0)

    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Add header text
    p_text.add_run(header_text)

    return table
