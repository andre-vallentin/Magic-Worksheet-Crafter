from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm

from .shared.table_column_width import set_column_width
from .shared.segment import Segment

class SingleChoiceTask(Segment):

    def __init__(self, title, description, content):
        super().__init__(title, description, content, icon_path=None)

    def build(self, doc):
        self.add_segment_header(doc)
        self.build_table(doc)
        doc.add_paragraph()  # Add spacing after the table        

    def build_table(self, doc):

        rows = len(self.content) + 1 # Add one for fixed header row
        cols = len(self.content[0])

        section = doc.sections[0]

        usable_width_emu = section.page_width - section.left_margin - section.right_margin
        usable_width_cm = usable_width_emu / 360000

        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.autofit = False

        table.cell(0, 0).text = 'Aussage'
        table.cell(0, 1).text = 'Richtig'
        table.cell(0, 2).text = 'Falsch'

        table.cell(0, 0).paragraphs[0].runs[0].bold = True
        table.cell(0, 1).paragraphs[0].runs[0].bold = True
        table.cell(0, 2).paragraphs[0].runs[0].bold = True

        table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(0, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set column widths using per-row cell approach
        set_column_width(table, 0, 10* (usable_width_cm - 3)) 
        set_column_width(table, 1, 15)   # 1.5 cm = 15 mm
        set_column_width(table, 2, 15)   # 1.5 cm = 15 mm

        for row_idx, row_data in enumerate(self.content):

            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx + 1, col_idx)  # +1 to skip header row    
                cell_text = str(cell_data)
                cell.text = cell_text
            table.cell(row_idx + 1, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for row in table.rows:
            row.height = Cm(1)

