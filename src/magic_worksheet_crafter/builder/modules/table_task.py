from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches

from .shared.segment import Segment

try:
    from pyphen import Pyphen
    HAS_PYPHEN = True
except ImportError:
    HAS_PYPHEN = False

class TableTask(Segment):
    # German hyphenator
    _hyphenator = None

    def __init__(self, title, description, content):
        super().__init__(title, description, content, icon_path=None)

    def build(self, doc):
        self.add_segment_header(doc)
        self.build_table(doc, self.content)        
        doc.add_paragraph()  # Add spacing after the table

    @classmethod
    def _get_hyphenator(cls):
        """Get or create German hyphenator."""
        if cls._hyphenator is None and HAS_PYPHEN:
            cls._hyphenator = Pyphen(lang='de_DE')
        return cls._hyphenator
    

    def build_table(self, doc, content):
        """Build a table from a 2D array.

        Args:
            doc: python-docx Document object
            content: 2D list where first row is headers, subsequent rows are data
        """
        if not content or not content[0]:
            return

        rows = len(content)
        cols = len(content[0])

        # Create table
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        # Calculate available cell width
        page_width = doc.sections[0].page_width
        left_margin = doc.sections[0].left_margin
        right_margin = doc.sections[0].right_margin
        table_width = page_width - left_margin - right_margin
        # Convert to EMU (English Metric Units), do arithmetic, then convert back to inches for twips
        # python-docx measures everything in EMU internally
        cell_width_emu = table_width // cols
        cell_width = Inches(cell_width_emu / 914400)  # 914400 EMU per inch

        # Fill table with data
        for row_idx, row_data in enumerate(content):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                # Apply German hyphenation based on cell width
                cell_text = str(cell_data)
                cell.text = cell_text

                # Enable word wrapping and allow hyphens
                tcPr = cell._element.get_or_add_tcPr()
                # Set cell properties to allow text wrapping
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(cell_width.twips)))  # Width in twips
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

                # Configure paragraph for word wrapping
                for paragraph in cell.paragraphs:
                    # Allow word wrapping with hyphenation
                    pPr = paragraph._element.get_or_add_pPr()
                    # Ensure paragraph is not set to keep together
                    keepNext = pPr.find(qn('w:keepNext'))
                    if keepNext is not None:
                        pPr.remove(keepNext)

                # Format header row
                if row_idx == 0:
                    # Set background color to light gray
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'e9e9e9')
                    cell._element.get_or_add_tcPr().append(shading_elm)

                    # Center alignment for header cells
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.runs[0].bold = True

                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Set row height to 1 cm for all rows
        for row in table.rows:
            row.height = Cm(1)