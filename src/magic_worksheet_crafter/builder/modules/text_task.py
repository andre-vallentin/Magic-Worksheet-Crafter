from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

from .shared.table_column_width import set_column_width
from .shared.markdown_handler import build_formatted_paragraph
from .shared.segment import Segment

class TextTask(Segment):

    def __init__(self, title, description, content):
        super().__init__(title, description, content, icon_path=None)
       
    def build(self, doc):
        self.add_segment_header(doc)
        self.add_half_linebreak(doc)
        
        for task in self.content:
            self.taskSection(doc, task)
            self.add_half_linebreak(doc)

        doc.add_paragraph()  # Add spacing after the table

    def add_half_linebreak(self, doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1.5)

    def taskSection(self, doc, task):
        build_formatted_paragraph(doc, task.description)

        if(task.answerLines == 0 and task.answerText != None):
            build_formatted_paragraph(doc, task.answerText)
            return

        table = doc.add_table(rows=task.answerLines, cols=1)
        table.autofit = False

        # Set explicit table width to prevent Word from ignoring it
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '0')
        tblW.set(qn('w:type'), 'auto')
        table._tbl.tblPr.insert(0, tblW)

        borders = OxmlElement('w:tblBorders')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '4')
        borders.append(bottom_border)
        table._tbl.tblPr.append(borders)

        # Set column width using per-row cell approach
        set_column_width(table, 0, 160)  # ~16cm

        for row in table.rows:
            row.height = Cm(1)
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()

                tcBorders = OxmlElement('w:tcBorders')
                bottom_border = OxmlElement('w:bottom')
                bottom_border.set(qn('w:val'), 'single')
                bottom_border.set(qn('w:sz'), '4')
                tcBorders.append(bottom_border)
                tcPr.append(tcBorders)                


    